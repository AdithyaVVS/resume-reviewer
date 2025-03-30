from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime
from resume_parser import ResumeParser

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your_secret_key_here')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resume_reviewer.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database
db = SQLAlchemy(app)

# Add context processor for templates
@app.context_processor
def inject_now():
    return {'now': datetime.now()}

# Initialize Resume Parser
try:
    resume_parser = ResumeParser()
    print("Successfully initialized Resume Parser")
except Exception as e:
    print(f"ERROR initializing Resume Parser: {str(e)}")
    resume_parser = None

# Database models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    resumes = db.relationship('Resume', backref='owner', lazy=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    analysis = db.Column(db.Text, nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

# Helper functions
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_resume(file_path):
    # Extract text based on file extension
    file_extension = file_path.rsplit('.', 1)[1].lower()
    
    if file_extension == 'pdf':
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            # Fallback to basic text extraction
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read().strip()
    
    elif file_extension == 'docx':
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs with proper spacing
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            # Fallback to basic text extraction
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read().strip()
    
    else:  # txt or other text formats
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read().strip()

def analyze_resume_with_ai(resume_text):
    try:
        if resume_parser is None:
            raise Exception("Resume parser not initialized. Check API key.")
        return resume_parser.parse_resume_text(resume_text)
    except Exception as e:
        print(f"Resume Analysis Error: {str(e)}")
        return f"Error: Unable to analyze resume. {str(e)}"

# Routes
@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        
        # Check if user already exists
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash('Username already exists')
            return redirect(url_for('register'))
        
        existing_email = User.query.filter_by(email=email).first()
        if existing_email:
            flash('Email already registered')
            return redirect(url_for('register'))
        
        # Create new user
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        flash('Registration successful! Please log in.')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            session['user_id'] = user.id
            flash('Login successful!')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    flash('You have been logged out')
    return redirect(url_for('index'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        flash('Please log in to access the dashboard')
        return redirect(url_for('login'))
    
    user = User.query.get(session['user_id'])
    resumes = Resume.query.filter_by(user_id=user.id).all()
    
    return render_template('dashboard.html', user=user, resumes=resumes)

@app.route('/upload', methods=['GET', 'POST'])
def upload_resume():
    if 'user_id' not in session:
        flash('Please log in to upload a resume')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Check if the post request has the file part
        if 'resume' not in request.files:
            flash('No file part')
            return redirect(request.url)
        
        file = request.files['resume']
        
        # If user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Extract text from resume
            try:
                resume_text = extract_text_from_resume(file_path)
                
                # Analyze resume using AI
                analysis = analyze_resume_with_ai(resume_text)
                
                # Serialize analysis data to JSON string before saving
                analysis_json = json.dumps(analysis)
                
                # Save resume and analysis to database
                new_resume = Resume(
                    filename=filename,
                    analysis=analysis_json,
                    user_id=session['user_id']
                )
                
                db.session.add(new_resume)
                db.session.commit()
                
                flash('Resume uploaded and analyzed successfully!')
                return redirect(url_for('view_analysis', resume_id=new_resume.id))
            
            except Exception as e:
                flash(f'Error analyzing resume: {str(e)}')
                return redirect(url_for('dashboard'))
    
    return render_template('upload.html')

@app.route('/analysis/<int:resume_id>')
def view_analysis(resume_id):
    if 'user_id' not in session:
        flash('Please log in to view analysis')
        return redirect(url_for('login'))
    
    resume = Resume.query.get_or_404(resume_id)
    
    # Check if the resume belongs to the logged-in user
    if resume.user_id != session['user_id']:
        flash('Unauthorized access')
        return redirect(url_for('dashboard'))
    
    # Check if analysis contains an error message
    if resume.analysis and resume.analysis.startswith('Error:'):
        flash(resume.analysis)
        return redirect(url_for('dashboard'))
    
    try:
        # Parse the stored JSON analysis data
        analysis_data = json.loads(resume.analysis) if resume.analysis else {}
        
        # Ensure all required sections exist with default values
        analysis = {
            'structure': {'strengths': [], 'improvements': [], 'suggestions': []},
            'skills': {'identified_skills': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'education': {'entries': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'experience': {'entries': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'ats': {'strengths': [], 'improvements': [], 'suggestions': []},
            'transformer_info': {'models_used': []}
        }
        
        # Update with actual data if available
        for section in analysis:
            if section in analysis_data:
                analysis[section].update(analysis_data[section])
    except (json.JSONDecodeError, TypeError) as e:
        print(f"Error parsing analysis data: {str(e)}")
        analysis = {
            'structure': {'strengths': [], 'improvements': [], 'suggestions': []},
            'skills': {'identified_skills': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'education': {'entries': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'experience': {'entries': [], 'strengths': [], 'improvements': [], 'suggestions': []},
            'ats': {'strengths': [], 'improvements': [], 'suggestions': []}
        }
    
    return render_template('analysis.html', resume=resume, analysis=analysis)

@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    if 'resume' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['resume']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            resume_text = extract_text_from_resume(file_path)
            analysis = analyze_resume_with_ai(resume_text)
            
            return jsonify({
                'success': True,
                'analysis': analysis
            })
        
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)